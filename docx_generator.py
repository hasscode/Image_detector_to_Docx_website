from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import re

def create_docx(clean_text: str, output_path: str):
    doc = Document()

    # 1. إعداد ستايل عام للخط
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # 2. الهيدر (Header) - لمسة المستشفى
    header = doc.sections[0].header
    htable = header.add_table(1, 2, doc.sections[0].page_width)
    htable.cell(0, 0).text = "Medical Radiology Department"
    htable.cell(0, 1).text = f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}"
    htable.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 3. العنوان الرئيسي
    title = doc.add_heading('FETAL BIOMETRY REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. معالجة النص وتحويل القياسات لجدول
    lines = clean_text.split('\n')
    biometry_data = [] # لتخزين القياسات زي BPD: 50mm
    
    for line in lines:
        line = line.strip()
        if not line: continue

        # البحث عن قياسات طبية (اسم : قيمة)
        match = re.match(r"^(BPD|HC|AC|FL|CRL|GA|EDD)\s*[:|-]\s*(.*)", line, re.IGNORECASE)
        
        if match:
            biometry_data.append(match.groups())
        elif line.startswith('#'):
            level = line.count('#')
            doc.add_heading(line.replace('#', '').strip(), level=min(level, 3))
        elif line.startswith('•') or line.startswith('-'):
            doc.add_paragraph(line.strip('•- '), style='List Bullet')
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 5. إذا وجدنا قياسات، نضعها في جدول منظم في نهاية التقرير أو مكانه
    if biometry_data:
        doc.add_heading('Fetal Biometry Summary Table', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Measurement'
        hdr_cells[1].text = 'Value / Result'
        
        for name, value in biometry_data:
            row_cells = table.add_row().cells
            row_cells[0].text = name.upper()
            row_cells[1].text = value

    # 6. الحفظ
    doc.save(output_path)